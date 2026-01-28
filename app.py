import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from pptx import Presentation
import io
import os
import nltk
import base64
import re
from collections import Counter
from fpdf import FPDF
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer


# --- 1. Robust NLTK Resource Management ---
def ensure_nltk_resources():
    """Ensures all necessary NLTK data is downloaded for analysis."""
    resources = ['punkt', 'punkt_tab', 'stopwords']
    for res in resources:
        try:
            nltk.data.find(f'tokenizers/{res}')
        except LookupError:
            nltk.download(res, quiet=True)


ensure_nltk_resources()

# --- 2. Page Configuration & Stealth UI ---
st.set_page_config(page_title="Professional AI Analyst", layout="wide")

# CSS to hide "Manage app" and clean up the interface
st.markdown("""
    <style>
    button[title="Manage app"] { display: none !important; }
    .stAppDeployButton { display: none !important; }
    footer {visibility: hidden;}
    [data-testid="stToolbar"] {visibility: hidden !important;}
    .block-container {padding-top: 1.5rem;}
    /* Enhance the look of the notes section */
    .stMarkdown h3 { color: #2E86C1; }
    </style>
    """, unsafe_allow_html=True)

# Sidebar for Settings
with st.sidebar:
    if os.path.exists("logo.png"):
        st.image("logo.png", use_container_width=True)
    st.title("Control Panel")
    st.markdown("---")
    analysis_mode = st.radio("Analysis Mode", ["Standard Review", "Executive Deep Dive"])
    # Increase sentence count for deep dive to capture "Full Work"
    sentence_count = 15 if analysis_mode == "Executive Deep Dive" else 8


# --- 3. AI Brain: Document Identification ---
def identify_document_type(text):
    """
    Analyzes text patterns to determine if it is a CV, Research,
    Lecture Notes, or Business Report.
    """
    text_lower = text.lower()

    # Keyword clusters
    cv_signals = ['experience', 'education', 'skills', 'objective', 'employment', 'curriculum vitae']
    research_signals = ['abstract', 'methodology', 'results', 'discussion', 'conclusion', 'references']
    notes_signals = ['lecture', 'chapter', 'module', 'topic', 'definition', 'introduction to']
    report_signals = ['executive summary', 'revenue', 'market', 'strategy', 'kpi', 'fiscal']

    scores = {
        "Professional CV / Resume": sum(1 for w in cv_signals if w in text_lower),
        "Academic Research Paper": sum(1 for w in research_signals if w in text_lower),
        "Educational / Lecture Notes": sum(1 for w in notes_signals if w in text_lower),
        "Strategic Business Report": sum(1 for w in report_signals if w in text_lower)
    }

    # Return the category with the highest score
    best_match = max(scores, key=scores.get)
    return best_match if scores[best_match] > 0 else "General Informational Document"


# --- 4. Systematic Reconstruction Logic ---
def systematic_rewrite(text, doc_type, summary_sentences):
    """
    Rewrites the document content into a professional, systematic format.
    Adds high-level insights and gap analysis.
    """
    # 1. Extract Intelligence (Keywords)
    words = re.findall(r'\w+', text.lower())
    # Filter for meaningful words (len > 5)
    common_words = [word for word, count in Counter(words).most_common(60) if len(word) > 4]
    top_keywords = common_words[:10]

    # 2. Build the Reconstructed Document
    # Header
    output = f"# 🏛️ SYSTEMATIC RECONSTRUCTION: {doc_type.upper()}\n\n"
    output += f"**Detected Core Topics:** {', '.join([f'`{k.upper()}`' for k in top_keywords[:6]])}\n\n"

    # Section A: The Systematic Rewrite (The "Full Work")
    output += "## 📑 1. Systematic Content Rewrite\n"
    output += "*The following is a structured reconstruction of the document's primary content, organized for high-level understanding:*\n\n"

    for i, s in enumerate(summary_sentences):
        sent_str = str(s)
        # Bold the key terms found in the text to show connectivity
        for k in top_keywords[:5]:
            sent_str = re.sub(f'({k})', r'**\1**', sent_str, flags=re.IGNORECASE)
        output += f"{i + 1}. {sent_str}\n\n"

    # Section B: Professional Enhancement (Filling Gaps)
    output += "## 💎 2. High-Level Professional Enhancement\n"
    output += f"**Context:** As a **{doc_type}**, this document requires specific sections to be considered 'Complete' by industry standards.\n\n"

    if "CV" in doc_type:
        output += "### ⚠️ Missing Strategic Elements:\n"
        output += f"1. **Quantifiable Impact:** The document lists responsibilities involving **{top_keywords[0]}**. It *must* instead list achievements (e.g., 'Increased efficiency by 20%').\n"
        output += "2. **Executive Summary:** A 3-sentence value proposition is missing at the top.\n"
    elif "Research" in doc_type:
        output += "### ⚠️ Academic Validity Check:\n"
        output += f"1. **Practical Application:** The theoretical data on **{top_keywords[0]}** is sound, but the document lacks a 'Real-World Implication' section.\n"
        output += "2. **Limitations:** No study is perfect. A 'Limitations of Study' paragraph should be added to increase credibility.\n"
    else:
        output += "### ⚠️ Professional Gaps:\n"
        output += f"1. **Actionable Conclusion:** The text explains **{top_keywords[0]}** well, but needs a specific 'Next Steps' or 'Recommendations' list.\n"
        output += "2. **Source Validation:** Ensure the data regarding **" + (
            top_keywords[1] if len(top_keywords) > 1 else "key topics") + "** is cited from updated sources.\n"

    # Section C: Lecture / Presentation Briefing
    output += "## 🔍 3. Presentation Strategy (Lecture Ready)\n"
    output += "If you are presenting this document, focus on this narrative flow:\n"
    output += f"> *\"This document explores the relationship between **{top_keywords[0]}** and **{top_keywords[1] if len(top_keywords) > 1 else 'outcome'}**. While the data suggests a strong correlation, we must critically examine the methodology used in section 2.\"*\n"

    return output


# --- 5. Highlighting Logic ---
def highlight_pdf_content(file_bytes, key_sentences):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page in doc:
        for sent in key_sentences:
            # We shorten the search string to 60 chars to improve match rate across line breaks
            text_instances = page.search_for(str(sent)[:60])
            for inst in text_instances:
                annot = page.add_highlight_annot(inst)
                annot.set_colors(stroke=(1, 1, 0))  # Yellow Highlighting
                annot.update()
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def highlight_docx_content(file_bytes, key_sentences):
    doc = Document(io.BytesIO(file_bytes))
    for para in doc.paragraphs:
        for sent in key_sentences:
            if str(sent)[:40] in para.text:
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# --- 6. Export Utilities ---
def create_pdf_export(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    # Sanitize text to avoid Latin-1 encoding errors in FPDF
    clean_text = text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 6, txt=clean_text)
    return pdf.output(dest='S').encode('latin-1')


# --- 7. Main Application Logic ---
st.title("🖋️ Professional AI Document Mirror")
st.write("Upload a document. The AI will Identify, Reconstruct, and Highlight it systematically.")

uploaded_file = st.file_uploader("Upload File (PDF or DOCX)", type=["pdf", "docx"])

if uploaded_file:
    file_bytes = uploaded_file.read()
    file_ext = uploaded_file.name.split(".")[-1].lower()

    with st.spinner("🤖 AI is reading, identifying, and rewriting your document..."):
        # A. Text Extraction
        raw_text = ""
        if file_ext == "pdf":
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                raw_text = " ".join([p.get_text() for p in doc])
        elif file_ext == "docx":
            doc = Document(io.BytesIO(file_bytes))
            raw_text = " ".join([p.text for p in doc.paragraphs])

        # B. Identification & Summarization
        doc_type = identify_document_type(raw_text)

        parser = PlaintextParser.from_string(raw_text, Tokenizer("english"))
        summarizer = LsaSummarizer()
        summary_sentences = summarizer(parser.document, sentence_count)

        # C. Systematic Reconstruction (The "Full Work")
        full_analysis_text = systematic_rewrite(raw_text, doc_type, summary_sentences)

        # D. Document Highlighting
        if file_ext == "pdf":
            processed_doc = highlight_pdf_content(file_bytes, summary_sentences)
            mime_type = "application/pdf"
        elif file_ext == "docx":
            processed_doc = highlight_docx_content(file_bytes, summary_sentences)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            processed_doc = file_bytes
            mime_type = "application/octet-stream"

    # --- Result Columns ---
    col_text, col_preview = st.columns([1, 1])

    with col_text:
        st.subheader(f"📝 Systematic Reconstruction ({doc_type})")
        st.markdown(full_analysis_text)

        st.divider()
        st.write("📥 **Download Professional Notes:**")
        b1, b2 = st.columns(2)
        b1.download_button("Download as Word", full_analysis_text, f"Reconstructed_{uploaded_file.name}.docx")
        b2.download_button("Download as PDF", create_pdf_export(full_analysis_text),
                           f"Reconstructed_{uploaded_file.name}.pdf")

    with col_preview:
        st.subheader("📄 Highlighted Preview")
        st.download_button(
            label=f"📥 Download Highlighted {file_ext.upper()}",
            data=processed_doc,
            file_name=f"Highlighted_{uploaded_file.name}",
            mime=mime_type,
            use_container_width=True
        )

        # Chrome Blocking Fix: Using <embed> instead of <iframe>
        if file_ext == "pdf":
            base64_pdf = base64.b64encode(processed_doc).decode('utf-8')
            pdf_display = f'<embed src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800px" type="application/pdf">'
            st.markdown(pdf_display, unsafe_allow_html=True)
        else:
            st.info("Preview is optimized for PDF. Please download the file to view highlights for Word.")

else:
    st.info("Please upload a document to begin the professional reconstruction.")